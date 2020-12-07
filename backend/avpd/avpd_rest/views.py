import numpy as np
from django.http import JsonResponse
from django.shortcuts import render

# Create your views here.
from django.views.decorators.csrf import csrf_exempt

from docx import Document
from rest_framework import status
from rest_framework.exceptions import ParseError
from rest_framework.parsers import FileUploadParser, JSONParser

# This is a really crappy way of doing this, if it is still here when you read this don't follow in my footsteps.
from rest_framework.response import Response
from rest_framework.views import APIView
from scipy.stats import norm

from avpd.avpd_rest.parsed_doc import ParsedDoc

essays = {}


class IDGenerator:
    def __init__(self):
        self._id = 0

    def newID(self):
        self._id += 1

        return str(self._id)


idGenerator = IDGenerator()


class EssaysView(APIView):
    parser_class = (FileUploadParser,)

    @csrf_exempt
    def post(self, request, format=None):
        if 'file' not in request.data:
            raise ParseError("Empty content")

        f = request.data['file']

        document = Document(f)

        text = ""
        for paragraph in document.paragraphs:
            text += paragraph.text

        essayID = idGenerator.newID()
        essays[essayID] = text

        response = Response(status=status.HTTP_201_CREATED)

        response['Location'] = '/users/steve_beve/essays/' + essayID

        return response


class EssayReportView(APIView):
    def get(self, request, pk, format=None):
        debug = essays
        for _, essay in essays.items():
            debug = ParsedDoc(essay)
        parsedEssays = {essayID: ParsedDoc(essay) for (essayID, essay) in essays.items()}
        otherEssays = [parsedEssay for (essayID, parsedEssay) in parsedEssays.items() if essayID != pk]

        essayData = np.array([[otherEssay.getPToWRatio(), otherEssay.getAverageWordLength()]
                              for otherEssay in otherEssays])

        essayDataMean = essayData.sum(axis=0) / len(essayData)
        essayDataSquaredMean = (essayData**2).sum(axis=0) / len(essayData)
        essayDataSD = np.sqrt(essayDataSquaredMean - essayDataMean**2)

        targetEssay = parsedEssays[pk]
        targetEssayData = np.array([targetEssay.getPToWRatio(), targetEssay.getAverageWordLength()])
        targetEssayStandardized = (targetEssayData - essayDataMean) / essayDataSD
        targetEssayABS = abs(targetEssayStandardized)

        pValues = 1. - (norm.cdf(targetEssayABS) - norm.cdf(-1 * targetEssayABS))
        probabilityOfAuthorship = pValues.prod()

        flag = False

        if probabilityOfAuthorship < 0.05:
            flag = True

        return JsonResponse({"probabilityOfAuthorship": pValues.prod(), "flag": flag})


# @csrf_exempt
# def essays(request):
#     if request.method == 'POST':
#         data = FileUploadParser().parse(request)
#         document = Document(data['file'])
#         text = ""
#         for paragraph in document.paragraphs:
#             text += paragraph.text
#
#         essayID = idGenerator.newID()
#         essays[essayID] = text
#
#         return JsonResponse({"id": essayID, "text": text})


@csrf_exempt
def test(request):
    if request.method == 'POST':
        variable = 5.
        return JsonResponse({"id": variable})



