from django.db import models
from docx import Document

from backend.api.models.essay import Essay
from backend.api.models.user import Instructor, Student
from notebooks import StyleProfile


class Classroom(models.Model):
    # Instructors can have many classrooms. Classrooms can have only one instructor.
    instructor = models.ForeignKey(Instructor, on_delete=models.CASCADE)
    # Classrooms can have many students. Students can be enrolled in many classrooms.
    students = models.ManyToManyField(Student)

    title = models.CharField(max_length=30)


class Assignment(models.Model):
    # Classrooms can have many assignments. Assignments can only belong to one classroom.
    classroom = models.ForeignKey(Classroom, on_delete=models.CASCADE)

    title = models.CharField(max_length=30)
    description = models.CharField(max_length=1000)
    due_date = models.DateTimeField()


class Submission(models.Model):
    # Students can post many submissions to an assignment. Submissions can only be posted to one assignment.
    assignment = models.ForeignKey(Assignment, on_delete=models.CASCADE)
    # Students can post many submissions. Each submission belongs to only one student.
    student = models.ForeignKey(Student, on_delete=models.CASCADE)

    date = models.DateTimeField()
    file = models.FileField()

    def contrast_report(self):
        """Generate the contrast report for this submission."""
        essays = Essay.objects.filter(student=self.student)
        documents = [Document(essay.file) for essay in essays]
        essays = [document_text(document) for document in documents]

        style_profile = StyleProfile()

        # Feed the student's history into their style profile.
        for essay in essays:
            style_profile.feed(essay)

        submission_essay = document_text(Document(self.file))

        # Score this submission based on the style profile.
        authorship_probability = style_profile.score(submission_essay)
        flag = authorship_probability < 0.05

        return {'authorship_probability': authorship_probability, 'flag': flag}


def document_text(document):
    text = ""

    for paragraph in document.paragraphs:
        text += paragraph.text

    return text