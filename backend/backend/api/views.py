import numpy as np
from django.http import JsonResponse
from django.shortcuts import render

# Create your views here.
from django.views.decorators.csrf import csrf_exempt

from docx import Document
from rest_framework import status
from rest_framework.decorators import api_view, permission_classes
from rest_framework.exceptions import ParseError
from rest_framework.parsers import FileUploadParser, JSONParser, MultiPartParser
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from rest_framework.views import APIView
from scipy.stats import norm

from backend.api.parsed_doc import ParsedDoc

from functools import wraps
import jwt

from django.http import JsonResponse


def get_token_auth_header(request):
    """Obtains the Access Token from the Authorization Header
    """
    auth = request.META.get("HTTP_AUTHORIZATION", None)
    parts = auth.split()
    token = parts[1]

    return token


def requires_scope(required_scope):
    """Determines if the required scope is present in the Access Token
    Args:
        required_scope (str): The scope required to access the resource
    """

    def require_scope(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            token = get_token_auth_header(args[0])
            decoded = jwt.decode(token, verify=False)
            if decoded.get("scope"):
                token_scopes = decoded["scope"].split()
                for token_scope in token_scopes:
                    if token_scope == required_scope:
                        return f(*args, **kwargs)
            response = JsonResponse({'message': 'You don\'t have access to this resource'})
            response.status_code = 403
            return response

        return decorated

    return require_scope


# These views are here to quickly check if the api is working.
@api_view(['GET'])
@permission_classes([AllowAny])
def public_dummy(request):
    return JsonResponse({'message': 'You are a public dummy!'})


@api_view(['GET'])
def private_dummy(request):
    return JsonResponse({'message': 'You are a dummy, but only in private.'})


# essays = {}
#
#
# class IDGenerator:
#     def __init__(self):
#         self._id = 0
#
#     def newID(self):
#         self._id += 1
#
#         return str(self._id)
#
#
# idGenerator = IDGenerator()
#
#
# class EssaysView(APIView):
#     parser_class = (FileUploadParser,)
#
#     @csrf_exempt
#     def post(self, request, format=None):
#         if 'file' not in request.data:
#             raise ParseError("Empty content")
#
#         f = request.data['file']
#
#         document = Document(f)
#
#         text = ""
#         for paragraph in document.paragraphs:
#             text += paragraph.text
#
#         essayID = idGenerator.newID()
#         essays[essayID] = text
#
#         response = Response(status=status.HTTP_201_CREATED, content_type="multipart/form-data")
#
#         response['Location'] = '/users/steve_beve/essays/' + essayID
#
#         return response
#
#
# class EssayReportView(APIView):
#     def get(self, request, pk, format=None):
#         debug = essays
#         for _, essay in essays.items():
#             debug = ParsedDoc(essay)
#         parsedEssays = {essayID: ParsedDoc(essay) for (essayID, essay) in essays.items()}
#         otherEssays = [parsedEssay for (essayID, parsedEssay) in parsedEssays.items() if essayID != pk]
#
#         essayData = np.array([[otherEssay.getPToWRatio(), otherEssay.getAverageWordLength()]
#                               for otherEssay in otherEssays])
#
#         essayDataMean = essayData.sum(axis=0) / len(essayData)
#         essayDataSquaredMean = (essayData ** 2).sum(axis=0) / len(essayData)
#         essayDataSD = np.sqrt(essayDataSquaredMean - essayDataMean ** 2)
#
#         targetEssay = parsedEssays[pk]
#         targetEssayData = np.array([targetEssay.getPToWRatio(), targetEssay.getAverageWordLength()])
#         targetEssayStandardized = (targetEssayData - essayDataMean) / essayDataSD
#         targetEssayABS = abs(targetEssayStandardized)
#
#         pValues = 1. - (norm.cdf(targetEssayABS) - norm.cdf(-1 * targetEssayABS))
#         probabilityOfAuthorship = pValues.prod()
#
#         flag = False
#
#         if probabilityOfAuthorship < 0.05:
#             flag = True
#
#         return JsonResponse({"probabilityOfAuthorship": pValues.prod(), "flag": flag})
#
#
# # @csrf_exempt
# # def essays(request):
# #     if request.method == 'POST':
# #         data = FileUploadParser().parse(request)
# #         document = Document(data['file'])
# #         text = ""
# #         for paragraph in document.paragraphs:
# #             text += paragraph.text
# #
# #         essayID = idGenerator.newID()
# #         essays[essayID] = text
# #
# #         return JsonResponse({"id": essayID, "text": text})
#
#
# @csrf_exempt
# def test(request):
#     if request.method == 'POST':
#         variable = 5.
#         return JsonResponse({"id": variable})
